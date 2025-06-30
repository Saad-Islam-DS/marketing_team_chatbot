import os
import re
import json
import streamlit as st
from agno.agent import Agent
from agno.tools.duckduckgo import DuckDuckGoTools
from agno.models.groq import Groq
from io import BytesIO
from docx import Document
import ast

# ------------------ Configuration ------------------
st.set_page_config(page_title="Hearst Crypto Studio", layout="wide")

try:
    os.environ["GROQ_API_KEY"] = st.secrets["GROQ_API_KEY"]
except Exception:
    st.error("GROQ_API_KEY not found in Streamlit secrets. Please add it to your Streamlit Cloud app secrets.")
    st.stop()

HEARST_WEBSITE = "https://www.hearstcorporation.io"
BRAND_COLORS = {
    "background": "#0A0F0F",
    "primary": "#2EFFAF",
    "text": "#FFFFFF",
    "secondary": "#1A2C38"
}

# Enhanced CSS for Hearst branding and modern look
st.markdown(f"""
<link href="https://fonts.googleapis.com/css2?family=Montserrat:wght@700;400&display=swap" rel="stylesheet">
<style>
    body, .stApp {{
        background: #000000 !important;
        color: #FFFFFF;
        font-family: 'Montserrat', Arial, sans-serif;
    }}
    .main-title {{
        font-family: 'Montserrat', Arial, sans-serif;
        font-size: 2.6rem;
        font-weight: 800;
        letter-spacing: 1.5px;
        margin-bottom: 0.5rem;
        text-shadow: 0 4px 24px #2EFFAF33;
        background: linear-gradient(90deg, #2EFFAF 0%, #00FFB2 50%, #00C6FB 100%);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        background-clip: text;
    }}
    .main-subtitle {{
        font-size: 1.2rem;
        color: #FFFFFFCC;
        margin-bottom: 2.5rem;
        font-weight: 400;
    }}
    .hearst-card {{
        background: rgba(26,44,56,0.98);
        border-radius: 16px;
        padding: 32px 36px 24px 36px;
        margin-bottom: 32px;
        box-shadow: 0 4px 24px rgba(46,255,175,0.10);
        border-left: 6px solid #2EFFAF;
        transition: box-shadow 0.2s;
    }}
    .hearst-card:hover {{
        box-shadow: 0 8px 32px rgba(46,255,175,0.18);
    }}
    .hearst-title {{
        color: #2EFFAF;
        font-size: 1.5rem;
        font-weight: 700;
        margin-bottom: 16px;
        display: flex;
        align-items: center;
        letter-spacing: 0.5px;
        text-shadow: 0 2px 8px rgba(46,255,175,0.10);
    }}
    .hearst-icon {{
        font-size: 2rem;
        margin-right: 14px;
        filter: drop-shadow(0 2px 8px #2EFFAF33);
    }}
    .hearst-links a {{
        color: #2EFFAF !important;
        text-decoration: underline !important;
        font-weight: 500;
        font-size: 1.05rem;
    }}
    .stTextInput input {{
        background-color: #1A2C38 !important;
        color: #2EFFAF !important;
        border: 2px solid #2EFFAF !important;
        border-radius: 6px;
        padding: 14px !important;
        font-size: 1.1rem;
    }}
    .stButton>button {{
        background: linear-gradient(90deg, #2EFFAF 0%, #00C6FB 100%) !important;
        color: #0A0F0F !important;
        border: none;
        font-weight: bold;
        padding: 14px 28px;
        border-radius: 6px;
        font-size: 1.1rem;
        box-shadow: 0 2px 8px rgba(46,255,175,0.10);
        transition: background 0.2s;
    }}
    .stButton>button:hover {{
        background: linear-gradient(90deg, #00C6FB 0%, #2EFFAF 100%) !important;
    }}
</style>
""", unsafe_allow_html=True)

# Improved main title and subtitle
st.markdown("""
<div style='text-align:center; margin-bottom: 2rem;'>
    <div class='main-title'>Hearst Marketing Team Helper</div>
    <div class='main-subtitle'>Empowering sustainable, data-driven crypto mining content for the future</div>
</div>
""", unsafe_allow_html=True)

st.markdown("<hr style='border:1px solid #2EFFAF; margin: 2rem 0;'>", unsafe_allow_html=True)

# ------------------ Agent Factories ------------------
@st.cache_resource
def create_research_agent():
    return Agent(
        name="Hearst Research Validator",
        role="Research the provided crypto mining topic and generate high-quality, data-driven content examples and a detailed research summary for the marketing team.",
        model=Groq(id="llama-3.3-70b-versatile"),
        tools=[DuckDuckGoTools(search=True, news=True)],
        instructions=[
            "You are a researcher for Hearst Corporation.",
            "Given a topic, perform in-depth research using only credible, high-quality industry sources (e.g., CoinDesk, Cointelegraph, PR Newswire, MiningWeek, MIT Tech Review, Cambridge Bitcoin Index, Academic papers from IEEE, Springer crypto conferences).",
            "Avoid social media and opinion blogs.",
            "Return ONLY valid JSON. Do NOT include any markdown, explanations, or text outside the JSON object.",
            "The JSON must have the following structure:",
            "",
            "{",
            '  "summary": "<1000-2000 word in-depth technical summary of the topic, including all major trends, findings, controversies, and future outlook. Use as much detail as possible, cite statistics, and reference sources.>",',
            '  "simple_explanation": "<A simple, clear explanation of the topic in 100-200 words, suitable for someone with no technical background.>",',
            '  "stats": ["statistic 1", "statistic 2", ...],',
            '  "links": ["Direct URL to a recent (last 12 months) research article or news post about the topic", "... (at least 8-10 unique, recent, high-quality links)"]',
            "}",
            "",
            "Example:",
            '{',
            '  "summary": "Bitcoin mining is the process of ... (very detailed, technical, long)",',
            '  "simple_explanation": "Bitcoin mining is how new bitcoins are created ... (simple, short)",',
            '  "stats": ["50% of bitcoin mining uses renewables", "Over 12,200 jobs provided by mining in Texas"],',
            '  "links": ["https://coindesk.com/2024/05/bitcoin-mining-sustainability", "https://cointelegraph.com/news/bitcoin-mining-2024-trends", "..."]',
            '}',
            "All links must be direct to the original research/news post, not homepages or aggregators, and as recent as possible."
        ],
        show_tool_calls=False,
        markdown=True
    )

@st.cache_resource
def create_post_agent():
    return Agent(
        name="Hearst Content Architect",
        role="Create data-driven, positive social posts for Hearst",
        model=Groq(id="llama-3.3-70b-versatile"),
        instructions=[
            "You are a content creator for Hearst Corporation.",
            "Given a research summary, statistics, and links, generate the following in JSON:",
            "",
            "{",
            '  "linkedin": "<LinkedIn post>",',
            '  "linkedin_reference": "<URL from the research links used in the post>",',
            '  "instagram": "<Instagram post>",',
            '  "instagram_reference": "<URL from the research links used in the post>",',
            '  "x": "<X (Twitter) post>",',
            '  "x_reference": "<URL from the research links used in the post>"',
            "}",
            "",
            "Each post should reference the research and statistics, but not include URLs in the post text. Provide the reference link separately, and always use a link from the provided research links."
        ],
        show_tool_calls=False,
        markdown=False
    )

# ------------------ Parsing Helpers ------------------
def parse_research(content: str):
    items = []
    blocks = re.split(r'### ', content)
    for block in blocks:
        if not block.strip():
            continue
        lines = block.strip().split('\n')
        entry = {"title": lines[0].strip()}
        body = '\n'.join(lines[1:])
        url_match = re.search(r'(http[s]?://[^\s]+)', body)
        if url_match:
            entry['url'] = url_match.group(1)
        date_match = re.search(r'Date[:]?\s*(.*)', body)
        if date_match:
            entry['date'] = date_match.group(1)
        stats = re.findall(r'(\d+\.?\d*%?\s*[^.;\n]+)', body)
        if stats:
            entry['stats'] = stats[:3]
        summary_match = re.search(r'(\d+\s-word technical summary[:]?.*)', body)
        # fallback: take 150 words after "Summary:" if present
        entry['summary'] = ''
        summary_lines = body.split('Summary:')
        if len(summary_lines) > 1:
            entry['summary'] = summary_lines[1].strip()
        items.append(entry)
    return items

def extract_json_from_text(text):
    """
    Extract the first JSON object found in a string, including from code blocks.
    """
    # Remove code block markers if present
    text = re.sub(r"```(?:json)?", "", text, flags=re.IGNORECASE).replace("```", "")
    match = re.search(r'\{[\s\S]*\}', text)
    if match:
        try:
            return json.loads(match.group(0))
        except json.JSONDecodeError:
            return None
    return None

def extract_json_objects(text):
    """
    Extract all JSON objects from a string, even if not in a JSON array.
    """
    json_objects = []
    # Find all {...} blocks that look like JSON objects
    for match in re.finditer(r'\{.*?\}', text, re.DOTALL):
        try:
            obj = json.loads(match.group(0))
            json_objects.append(obj)
        except Exception:
            continue
    return json_objects

# ------------------ Display Helpers ------------------
def display_research(items):
    st.subheader("🔍 Verified Research Breakdown")
    for idx, item in enumerate(items, 1):
        with st.expander(f"Research #{idx}: {item.get('title', 'No Title')}", expanded=False):
            if 'url' in item:
                st.markdown(f"**Source:** [{item['url']}]({item['url']})")
            if 'date' in item:
                st.markdown(f"**Date:** {item['date']}")
            if 'stats' in item:
                st.markdown("**Key Statistics:**")
                for stat in item['stats']:
                    st.markdown(f"- {stat}")
            if item.get('summary'):
                st.markdown("**Technical Summary:**")
                st.markdown(item['summary'])

# Display posts in cards
def display_hearst_post(title, icon, content):
    st.markdown(f"""
    <div class="hearst-card">
        <div class="hearst-title"><span class="hearst-icon">{icon}</span>{title}</div>
        <div>{content}</div>
    </div>
    """, unsafe_allow_html=True)


# Word File Creation
def create_word_doc(posts, research_data):
    from docx import Document
    from io import BytesIO
    doc = Document()
    doc.add_heading('Hearst Crypto Marketing Suite', 0)

    # Posts
    for plat in ['linkedin', 'instagram', 'x']:
        doc.add_heading(f"{plat.title()} Post", level=1)
        doc.add_paragraph(posts.get(plat, 'No post found.'))

    # Full Technical Research Summary
    doc.add_heading("Full Technical Research Summary", level=1)
    doc.add_paragraph(research_data.get('summary', 'No research summary found.'))

    # Simple Explanation
    doc.add_heading("Simple Explanation", level=1)
    doc.add_paragraph(research_data.get('simple_explanation', 'No simple explanation found.'))

    # Research Links
    doc.add_heading("Research Links", level=1)
    for url in research_data.get('links', []):
        doc.add_paragraph(url, style='List Bullet')

    # Save to BytesIO
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def create_ideas_word_doc(ideas):
    from docx import Document
    from io import BytesIO
    doc = Document()
    doc.add_heading('Hearst Crypto Mining Ideas', 0)
    for idx, idea in enumerate(ideas, 1):
        doc.add_heading(f"{idx}. {idea.get('topic', 'No Topic')}", level=1)
        doc.add_paragraph(f"Summary: {idea.get('summary', '')}")
        doc.add_paragraph("Useful Links:")
        for url in idea.get('links', []):
            doc.add_paragraph(url, style='List Bullet')
        doc.add_paragraph(f"Description: {idea.get('description', '')}")
        if idea.get('suggested_post_angle'):
            doc.add_paragraph(f"Suggested Post Angle: {idea.get('suggested_post_angle', '')}")
        doc.add_paragraph("")  # Spacer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def parse_ideas_from_text(text):
    # Split by idea sections (look for ### or numbered headings)
    idea_blocks = re.split(r'(?:^|\n)#+\s*\d+\.\s*|(?:^|\n)(?=\d+\.\s)', text)
    ideas = []
    for block in idea_blocks:
        block = block.strip()
        if not block or len(block) < 20:
            continue
        idea = {}
        # Topic
        topic_match = re.search(r'\*\*Topic:\*\*\s*(.*)', block)
        if not topic_match:
            topic_match = re.search(r'Topic:\s*(.*)', block)
        idea['topic'] = topic_match.group(1).strip() if topic_match else block.split('\n')[0].strip()
        # Summary
        summary_match = re.search(r'\*\*Summary:\*\*\s*(.*)', block)
        if not summary_match:
            summary_match = re.search(r'Summary:\s*(.*)', block)
        idea['summary'] = summary_match.group(1).strip() if summary_match else ""
        # Links
        links = re.findall(r'\[([^\]]+)\]\((https?://[^\)]+)\)', block)
        if not links:
            # fallback: plain links
            links = [(url, url) for url in re.findall(r'(https?://[^\s\*]+)', block)]
        idea['links'] = [url for _, url in links]
        # Description
        desc_match = re.search(r'\*\*Description:\*\*\s*(.*)', block, re.DOTALL)
        if not desc_match:
            desc_match = re.search(r'Description:\s*(.*)', block, re.DOTALL)
        desc = desc_match.group(1).strip() if desc_match else ""
        # Remove "Suggested Post Angle" from description if present
        desc = re.split(r'\*\*Suggested Post Angle:\*\*|Suggested Post Angle:', desc)[0].strip()
        idea['description'] = desc
        # Suggested Post Angle
        angle_match = re.search(r'\*\*Suggested Post Angle:\*\*\s*"(.*?)"', block)
        if not angle_match:
            angle_match = re.search(r'Suggested Post Angle:\s*"(.*?)"', block)
        idea['suggested_post_angle'] = angle_match.group(1).strip() if angle_match else ""
        ideas.append(idea)
    return ideas

def is_valid_article_link(url):
    # Basic filter: must not end with /, /topic/, /news/, etc.
    invalid_patterns = ['/topic/', '/news/', '/tags/', '/category/', '/search?', '/?']
    if any(p in url for p in invalid_patterns):
        return False
    # Must look like an article (has a date or slug)
    return re.search(r'/\d{4}/\d{2}/\d{2}/|/article|/news|/story|/202', url) is not None

# ------------------ App UI ------------------


tab1, tab2 = st.tabs(["📢 Post Generator", "💡 Idea Lab"])

with tab1:
    st.header("Professional Content Creation")
    topic = st.text_input(
        "Enter Bitcoin mining topic:",
        placeholder="e.g., 'Renewable energy in Bitcoin mining'"
    )
    if st.button("Generate Expert Content"):
        with st.spinner("Researching and generating content..."):
            # Phase 1: Research
            research_agent = create_research_agent()
            research_raw = research_agent.run(topic)

            # Try direct JSON load first, then fallback to extraction
            try:
                research_data = json.loads(research_raw.content)
            except json.JSONDecodeError:
                research_data = extract_json_from_text(research_raw.content)
                if not research_data:
                    st.error("Failed to parse research output. Received: " + research_raw.content)
                    st.stop()

            if not research_data:
                st.warning("No valid research found. Try broadening your topic or adjusting keywords.")
            else:
                # Phase 2: Post Generation
                post_agent = create_post_agent()
                post_input = {
                    "summary": research_data.get("summary", ""),
                    "stats": research_data.get("stats", []),
                    "links": research_data.get("links", [])
                }
                post_raw = post_agent.run(json.dumps(post_input))
                try:
                    posts = json.loads(post_raw.content)
                except json.JSONDecodeError:
                    posts = extract_json_from_text(post_raw.content)
                    if not posts:
                        st.error("Failed to parse posts output. Received: " + post_raw.content)
                        st.stop()

                # 1. LinkedIn Post Example
                display_hearst_post(
                    "LinkedIn Post Example", "💼",
                    posts.get('linkedin', 'No LinkedIn post found.') +
                    (f"<br><span style='color:#2EFFAF;font-size:0.95rem;'>Reference: <a href='{posts.get('linkedin_reference','')}' target='_blank'>{posts.get('linkedin_reference','')}</a></span>" if posts.get('linkedin_reference') else "")
                )

                # 2. Instagram Post Example
                display_hearst_post(
                    "Instagram Post Example", "📸",
                    posts.get('instagram', 'No Instagram post found.') +
                    (f"<br><span style='color:#2EFFAF;font-size:0.95rem;'>Reference: <a href='{posts.get('instagram_reference','')}' target='_blank'>{posts.get('instagram_reference','')}</a></span>" if posts.get('instagram_reference') else "")
                )

                # 3. X Post Example
                display_hearst_post(
                    "X Post Example", "✖️",
                    posts.get('x', 'No X post found.') +
                    (f"<br><span style='color:#2EFFAF;font-size:0.95rem;'>Reference: <a href='{posts.get('x_reference','')}' target='_blank'>{posts.get('x_reference','')}</a></span>" if posts.get('x_reference') else "")
                )

                # 4. Full Technical Research Summary
                display_hearst_post("Full Technical Research Summary", "📑", research_data.get('summary', 'No research summary found.'))

                # 5. Simple Explanation
                display_hearst_post("Simple Explanation", "📝", research_data.get('simple_explanation', 'No simple explanation found.'))

                # 6. Research Links
                links_html = "<ul class='hearst-links'>" + "".join(
                    f"<li><a href='{url}' target='_blank'>{url}</a></li>" for url in research_data.get('links', [])
                ) + "</ul>"
                display_hearst_post("Research Links", "🔗", links_html)

        # Export option
        if 'posts' in locals():
            word_buffer = create_word_doc(posts, research_data)
            st.download_button(
                label="⬇️ Download Word Report",
                data=word_buffer,
                file_name="hearst_posts.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key="download_word"
            )

with tab2:
    st.header("Strategic Idea Generation")
    if st.button("Generate 5 Expert Ideas"):
        with st.spinner("Gathering latest mining news..."):
            # Agent 1: Gather and summarize latest mining news
            mining_news_agent = Agent(
                name="Crypto Mining News Summarizer",
                role="Summarize the latest (last 1-2 months) crypto mining news and provide a list of direct article links.",
                model=Groq(id="llama-3.3-70b-versatile"),
                tools=[DuckDuckGoTools(search=True, news=True)],
                instructions=[
                    "Research the latest (last 1-2 months) news and innovations in the crypto mining industry ONLY and make sure to give urls of the research.",
                    "Return a JSON object with:",
                    '{ "summary": "<detailed summary of the latest news>", "links": ["direct article link 1", "direct article link 2", "... (at least 5)"] }',
                    "All links must be direct to a specific article, not homepages or topic pages."
                ],
                show_tool_calls=False,
                markdown=False
            )
            mining_news = mining_news_agent.run("Latest crypto mining news and innovations")
            try:
                mining_news_data = json.loads(mining_news.content)
            except Exception:
                mining_news_data = extract_json_from_text(mining_news.content)
            if not mining_news_data or not mining_news_data.get("links"):
                st.error("Failed to get mining news. Try again.")
                st.stop()

            # Agent 2: Generate 5 ideas from the summary and links
            idea_agent = Agent(
                name="Crypto Mining Idea Generator",
                role="Generate actionable content ideas for Hearst based on mining news.",
                model=Groq(id="llama-3.3-70b-versatile"),
                instructions=[
                    "Given the following mining news summary and links, generate 5 actionable content ideas for Hearst Corporation.",
                    "Each idea must be about crypto mining or a new technology directly impacting mining.",
                    "Return a JSON array of 5 objects, each with:",
                    '{ "topic": "<short, catchy idea title>", "summary": "<5-10 sentence summary>", "description": "<detailed explanation>", "links": ["useful article link(s) from the provided list"], "suggested_post_angle": "<optional>" }'
                ],
                show_tool_calls=False,
                markdown=False
            )
            idea_input = {
                "summary": mining_news_data.get("summary", ""),
                "links": mining_news_data.get("links", [])
            }
            ideas = idea_agent.run(json.dumps(idea_input))
            try:
                ideas_list = json.loads(ideas.content)
            except Exception:
                ideas_list = extract_json_objects(ideas.content)
            if not ideas_list:
                st.error("Failed to parse ideas output. Received: " + ideas.content)
                st.stop()

            # Display ideas in a simple vertical stack (VBox)
            for idx, idea in enumerate(ideas_list, 1):
                st.markdown(f"""
                <div class="hearst-card">
                    <div class="hearst-title"><span class="hearst-icon">💡</span>{idx}. {idea.get('topic', 'No Topic')}</div>
                    <div><b>Summary:</b> {idea.get('summary', '')}</div>
                    <div style="margin-top:0.5em;"><b>Description:</b> {idea.get('description', '')}</div>
                    {f"<div style='margin-top:0.5em;'><b>Useful Links:</b><ul class='hearst-links'>{''.join(f'<li><a href={url} target=_blank>{url}</a></li>' for url in idea.get('links', []))}</ul></div>" if idea.get('links') else ''}
                    {f"<div style='margin-top:0.5em;'><b>Suggested Post Angle:</b> {idea.get('suggested_post_angle', '')}</div>" if idea.get('suggested_post_angle') else ""}
                </div>
                """, unsafe_allow_html=True)

            # Download as Word
            if 'ideas_list' in locals():
                ideas_word_buffer = create_ideas_word_doc(ideas_list)
                st.markdown("<div style='margin-top:2rem; text-align:center;'><b style='color:#2EFFAF; font-size:1.15rem;'>Download these ideas as a Word document:</b></div>", unsafe_allow_html=True)
                st.download_button(
                    label="⬇️ Download Ideas Word Report",
                    data=ideas_word_buffer,
                    file_name="hearst_ideas.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key="download_ideas_word"
                )

